import json
from logging import warning
import os
import re
from docx import Document
from win32com import client as wc
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice 
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
import jieba 
import math
import random

class Data_prepossed():
    def read_file(self,filename):#读取数据
        data=[]
        with open(filename,'r',encoding='utf-8') as file:
            for line in file.readlines():
                line=line.strip()
                data.append(line)
        return data
    def read_docx(self,filename):#读取docx内容
        docu=Document(filename)
        data=[]
        for para in docu.paragraphs:
            if para.text!='':
                data.append(para.text)
        return data
    #读取pdf内容
    def read_pdf(self,filename):
        print(filename)
        data=[]
        fp=open(filename,'rb')
        #创建解释器
        pdf_parser=PDFParser(fp)
        #pdf文档对象
        doc=PDFDocument()
        #连接解释器与文档对象
        pdf_parser.set_document(doc)
        doc.set_parser(pdf_parser)
        #初始化文档
        doc.initialize()
        #创建PDF资源管理器
        resource=PDFResourceManager()
        #创建PDF参数分析器
        laparam=LAParams()
        #创建聚合器
        device = PDFPageAggregator(resource, laparams=laparam)
 
        #创建PDF页面解析器
        interpreter = PDFPageInterpreter(resource, device)
        for page in doc.get_pages():
            #使用页面解释器来读取
            interpreter.process_page(page)
            #使用聚合器获得内容
            layout = device.get_result()
            for out in layout:       
                if hasattr(out, 'get_text'):#判断对象是否包含get_text()属性
                    if out.get_text()!='':
                        data.append(out.get_text().strip())
        fp.close()
        return data

    #将doc文件批量转换为docx文件
    def convert_docx(self,filename):
        split_words=filename.split('.')
        filetype=split_words[1]
        name=split_words[0]
        if filetype=='doc':
            word = wc.Dispatch('Word.Application')#启动进程
            doc = word.Documents.Open(filename)
            doc.SaveAs(name+'.docx', 12, False, "", True, "", False, False, False, False)  #转化后路径下的文件    
            doc.Close()
            word.Quit()
    def LoadStopWords(self,filepath):#加载停用词表
        WordsList=[]
        with open(filepath,'r',encoding='utf-8') as file:
            for line in file.readlines():
                line=line.strip('\n')
                WordsList.append(line)
        return WordsList
    #对每一个句子进行分词
    def segmentWords(self,sentence):
        WordsList=self.LoadStopWords('./stopwords(new).txt')
        sentence=''.join(sentence.split())#去除空格
        seg_words=[x for x in jieba.cut(sentence)]#使用jieba库分词精确模式，得到单词列表
        seg_words=list(filter(lambda x: x not in WordsList,seg_words))
        return seg_words
    #得到附件的文件内容,并将其写入到json文件中
    def get_file_data(self,filepath):
        file_data=[]#key为文件名，value为列表，由于文件读取的原因，列表中的元素是一串句子或者一段或者是短语
        label=['filename','content','level']
        result_dic=[]
        for filenames in os.walk(filepath):
            dirpath=filenames[0]
            filenames=filenames[2]
            for i in range(len(filenames)):
                filename=filepath+'\\'+filenames[i]
                filetype=filename.split('.')[1]
                '''
                if filetype=='pdf':
                    file_data.setdefault(filenames[i],self.read_pdf(filename))
                '''
                if filetype=='docx':
                    level=random.randint(1,4)
                    result=self.read_docx(filename)
                    content=''
                    for j in range(len(result)):
                        content+=(result[j].strip())
                    file_data.append([filenames[i],content,level])
        for i in range(len(file_data)):
            temp_dic=dict(zip(label,file_data[i]))
            str_json=json.dumps(temp_dic,ensure_ascii=False)
            result_dic.append(str_json)

        with open('file_data.json','w',encoding='utf-8') as wfile:
            for i in range(len(result_dic)):
                wfile.write(result_dic[i]+'\n')
        return 
    def get_web_data(self,filename):
        web_data=[]#列表，每一个列表元素也为一个列表
        data=self.read_file(filename)
        for i in range(len(data)):
            level=random.randint(1,4)
            temp_dic=json.loads(data[i])
            web_data.append([temp_dic['url'],temp_dic['content'],temp_dic['title'],level])
        return web_data
    def create_web_index(self,query,userlever):
        web_data=self.get_web_data('./web_data.json')
        tf={}#统计每个页面下词频,相当于倒排索引
        df={}#统计单词出现的文本数量
        titleset={}
        urlset=[]
        #计算文本平均长度
        web_len={}
        for i in range(len(web_data)):
            urlset.append([web_data[i][0],web_data[i][3]])#网址和权限集合
            web_len.setdefault(web_data[i][0],len(web_data[i][1]))
            titleset.setdefault(web_data[i][0],web_data[i][2])
        avg_len=0
        for value in web_len.values():
            avg_len+=value
        avg_len=round(avg_len/len(web_len.values()),6)
        #统计词频和单词出现的文本数量
        all_list=[]
        for i in range(len(web_data)):
            all_list.append([self.segmentWords(web_data[i][1]),web_data[i][0]])
        self.cal_nums(all_list,df,tf)
        query_result=self.cal_bm(tf,df,query,urlset,avg_len,web_len,userlever)
        result=''
        for i in range(len(query_result)):
            result+=str(query_result[i])+'\t'+str(titleset[query_result[i]])+'\n'
        return result
    
    #创建附件的索引
    def create_file_index(self,filepath,query,userlever):
        data=self.read_file(filepath)
        file_data=[]
        for i in range(len(data)):
            temp_dic=json.loads(data[i])
            file_data.append([temp_dic['filename'],temp_dic['content'],temp_dic['level']])
        tf={}#统计每个文档词频,相当于倒排索引
        df={}#统计单词出现的文本数量
        fileset=[]
        #计算文本平均长度
        file_len={}
        for i in range(len(file_data)):
            fileset.append([file_data[i][0],file_data[i][2]])#文件名和权限
            file_len.setdefault(file_data[i][0],len(file_data[i][1]))
        avg_len=0
        for value in file_len.values():
            avg_len+=value
        avg_len=round(avg_len/len(file_len.values()),6)
        #统计词频和单词出现的文本数量
        all_list=[]
        for i in range(len(file_data)):
            all_list.append([self.segmentWords(file_data[i][1]),file_data[i][0]])
        self.cal_nums(all_list,df,tf)
        query_result=self.cal_bm(tf,df,query,fileset,avg_len,file_len,userlever)
        result=''
        for i in range(len(query_result)):
            result+=str(query_result[i])+'\n'
        return result
    

    #统计词频和单词出现的文章数目
    def cal_nums(self,all_list,df,tf):
        for i in range(len(all_list)):
            words_list=all_list[i][0]
            url=all_list[i][1]
            for word in words_list:
                if word not in tf.keys():#该单词第一次出现
                    tf.setdefault(word,{url:1})
                    if word not in df.keys():
                        df.setdefault(word,1)
                    else:
                        df[word]+=1
                else:
                    if url not in tf[word].keys():#该文章该单词第一次出现
                        tf[word].setdefault(url,1)
                        df[word]+=1
                    else:
                        tf[word][url]+=1
    #计算BM2.5
    def cal_bm(self,tf,df,query,urlset,avglen,web_len,userlevel):#query是问题，需要对其进行分词，level表示用户的权限
        query=self.segmentWords(query)
        score={}
        N=len(urlset)
        #计算每一篇文档和查询的相关性，用bm2.5表示
        for m in range(len(urlset)):
            url=urlset[m][0]
            level=urlset[m][1]
            if level>userlevel:#用户权限低于要查询内容的权限
                continue
            temp_score=0
            K=2*(1-0.75+0.75*(web_len[url]/avglen))
            for i in range(len(query)):
                query_word=query[i]
                if query_word not in tf.keys() or url not in tf[query_word].keys():#查询中的词未在所有的文本中出现
                    continue
                wi=math.log((N-df[query_word]+0.5)/(df[query_word]+0.5))
                fi=tf[query_word][url]
                Ri=3*fi/(fi+K)
                temp_score+=wi*Ri
            score.setdefault(url,temp_score)
        a=sorted(score.items(),key=lambda x:x[1],reverse=True)
        result=[]
        for i in range(min(5,len(a))):
            result.append(a[i][0])
        return result

if __name__=='__main__':
    Reader=Data_prepossed()
    filepath='F:\\pythoncode\\IRLab3\\Appendix'
    query='坚决维护习近平总书记党中央的核心'
    result=Reader.create_web_index(query,4)
    print(result)
    file_reasult=Reader.create_file_index('./file_data.json',query,1)    
    print(file_reasult)
    #Reader.get_file_data(filepath)#读取文件内容
    
